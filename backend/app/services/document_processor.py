import os
import io
import logging
from typing import List, Dict, Any, Optional
from PyPDF2 import PdfReader
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from app.core.config import settings

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Service for processing and chunking documents."""
    
    def __init__(self):
        """Initialize document processor with text splitter."""
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""]
        )
    
    def extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file."""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PdfReader(pdf_file)
            text = ""
            
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += f"\n[Page {page_num + 1}]\n{page_text}"
            
            return text
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise
    
    def extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file."""
        try:
            doc_file = io.BytesIO(file_content)
            doc = docx.Document(doc_file)
            text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text += cell.text + " "
                    text += "\n"
            
            return text
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise
    
    def extract_text_from_txt(self, file_content: bytes) -> str:
        """Extract text from TXT file."""
        try:
            return file_content.decode('utf-8', errors='ignore')
        except Exception as e:
            logger.error(f"Error extracting text from TXT: {str(e)}")
            raise
    
    def extract_text(self, file_content: bytes, file_type: str) -> str:
        """
        Extract text from various file types.
        
        Args:
            file_content: File content as bytes
            file_type: File extension (pdf, docx, txt, etc.)
        
        Returns:
            Extracted text
        """
        file_type = file_type.lower().replace('.', '')
        
        if file_type == 'pdf':
            return self.extract_text_from_pdf(file_content)
        elif file_type in ['docx', 'doc']:
            return self.extract_text_from_docx(file_content)
        elif file_type in ['txt', 'text']:
            return self.extract_text_from_txt(file_content)
        else:
            # Default to treating as text
            return self.extract_text_from_txt(file_content)
    
    def chunk_document(self, text: str, metadata: Optional[Dict] = None) -> List[Dict[str, Any]]:
        """
        Chunk document text into smaller pieces.
        
        Args:
            text: Document text
            metadata: Optional metadata to attach to chunks
        
        Returns:
            List of chunks with metadata
        """
        try:
            # Split text into chunks
            chunks = self.text_splitter.split_text(text)
            
            # Format chunks with metadata
            formatted_chunks = []
            for i, chunk_text in enumerate(chunks):
                chunk_data = {
                    'text': chunk_text,
                    'chunk_index': i,
                    'total_chunks': len(chunks)
                }
                
                # Add any additional metadata
                if metadata:
                    chunk_data.update(metadata)
                
                formatted_chunks.append(chunk_data)
            
            logger.info(f"Created {len(formatted_chunks)} chunks from document")
            return formatted_chunks
            
        except Exception as e:
            logger.error(f"Error chunking document: {str(e)}")
            raise
    
    def process_document(
        self, 
        file_content: bytes, 
        file_name: str,
        file_type: str,
        metadata: Optional[Dict] = None
    ) -> Dict[str, Any]:
        """
        Process a complete document: extract text and chunk it.
        
        Args:
            file_content: File content as bytes
            file_name: Name of the file
            file_type: File extension
            metadata: Optional metadata
        
        Returns:
            Processed document with chunks
        """
        try:
            # Extract text
            text = self.extract_text(file_content, file_type)
            
            if not text or len(text.strip()) == 0:
                raise ValueError("No text could be extracted from the document")
            
            # Prepare metadata
            doc_metadata = {
                'file_name': file_name,
                'file_type': file_type,
                'total_characters': len(text),
                **(metadata or {})
            }
            
            # Chunk the document
            chunks = self.chunk_document(text, {'file_name': file_name})
            
            return {
                'text': text,
                'chunks': chunks,
                'metadata': doc_metadata,
                'num_chunks': len(chunks)
            }
            
        except Exception as e:
            logger.error(f"Error processing document {file_name}: {str(e)}")
            raise
    
    def extract_key_information(self, text: str) -> Dict[str, Any]:
        """
        Extract key information from legal documents.
        This is a placeholder for more sophisticated NER/extraction.
        
        Args:
            text: Document text
        
        Returns:
            Extracted entities and key information
        """
        # This would typically use NER models or regex patterns
        # For now, returning placeholder structure
        return {
            'parties': [],
            'dates': [],
            'monetary_amounts': [],
            'legal_terms': [],
            'obligations': [],
            'signatures': []
        }

# Singleton instance
document_processor = None

def get_document_processor() -> DocumentProcessor:
    """Get or create document processor instance."""
    global document_processor
    if document_processor is None:
        document_processor = DocumentProcessor()
    return document_processor 